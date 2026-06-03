# Generates the detailed Word (.docx) Course Creation admin guide.
# Run:  python scripts/gen-course-creation-guide-docx.py
# Output: ../The-Raising-Club-Course-Creation-Guide.docx  (project root, not in git)

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK     = RGBColor(0x2B, 0x2B, 0x2B)
PRIMARY = RGBColor(0xC0, 0x53, 0x3B)
GREEN   = RGBColor(0x5E, 0x7B, 0x3A)
MUTED   = RGBColor(0x6B, 0x6B, 0x6B)
AMBER   = RGBColor(0x9A, 0x6B, 0x00)
PURPLE  = RGBColor(0x6B, 0x4E, 0x8C)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def h1(text, before=18):
    p = doc.add_heading(text, level=1)
    for r in p.runs: r.font.color.rgb = PRIMARY; r.font.size = Pt(18)
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(6)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs: r.font.color.rgb = INK; r.font.size = Pt(13.5)
    p.paragraph_format.space_before = Pt(12)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    for r in p.runs: r.font.color.rgb = GREEN; r.font.size = Pt(11.5)
    p.paragraph_format.space_before = Pt(8)
    return p

def para(text="", bold=False, italic=False, color=None, size=None, after=6):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.bold = bold; r.italic = italic
    if color is not None: r.font.color.rgb = color
    if size is not None:  r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(after)
    return p

def rich(parts, after=6):
    p = doc.add_paragraph()
    for t in parts:
        text = t[0]; bold = t[1] if len(t) > 1 else False; italic = t[2] if len(t) > 2 else False
        r = p.add_run(text); r.bold = bold; r.italic = italic
    p.paragraph_format.space_after = Pt(after)
    return p

def bullet(text, sub=False):
    p = doc.add_paragraph(style="List Bullet" + (" 2" if sub else ""))
    p.add_run(text); p.paragraph_format.space_after = Pt(2)
    return p

def steps(items):
    for i, it in enumerate(items, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}.  "); run.bold = True; run.font.color.rgb = PRIMARY
        p.add_run(it)
        p.paragraph_format.left_indent = Inches(0.22); p.paragraph_format.space_after = Pt(3)

def callout(label, text, color=AMBER):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"{label}  "); r.bold = True; r.font.color.rgb = color
    r2 = p.add_run(text); r2.italic = True; r2.font.color.rgb = MUTED

def field_table(rows, headers=("Field", "What it is / what to enter")):
    t = doc.add_table(rows=1, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        rp = hdr[i].paragraphs[0].add_run(htext); rp.bold = True
        rp.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); rp.font.size = Pt(10)
        _shade(hdr[i], "C0533B")
    for name, desc in rows:
        cells = t.add_row().cells
        cells[0].text = ""; r0 = cells[0].paragraphs[0].add_run(name); r0.bold = True; r0.font.size = Pt(10)
        cells[1].text = ""; cells[1].paragraphs[0].add_run(desc).font.size = Pt(10)
    for row in t.rows:
        row.cells[0].width = Inches(1.9); row.cells[1].width = Inches(4.7)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def page_break(): doc.add_page_break()

# ===========================================================================
# COVER
# ===========================================================================
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("The Raising Club"); tr.font.size = Pt(30); tr.bold = True; tr.font.color.rgb = PRIMARY
title.paragraph_format.space_before = Pt(120)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Course Creation — Admin Guide"); sr.font.size = Pt(20); sr.font.color.rgb = INK
sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
s2 = sub2.add_run("Build, price, publish and manage training courses — every screen, field and option explained")
s2.italic = True; s2.font.color.rgb = MUTED; s2.font.size = Pt(11)
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(40)
mr = meta.add_run("For new administrators · No technical background required"); mr.font.color.rgb = MUTED; mr.font.size = Pt(10)
page_break()

# ===========================================================================
# CONTENTS
# ===========================================================================
h1("What's in this guide", before=0)
for c in [
    "1.   Before you begin — what a course is, and how learners experience it",
    "2.   The Courses dashboard",
    "3.   The course editor at a glance (4 tabs + the save bar)",
    "4.   Tab 1 — Course details (and Skills)",
    "5.   Tab 2 — Chapters & modules",
    "6.   Tab 3 — Integration Moment (the quiz)",
    "7.   Tab 4 — Certificate",
    "8.   Pricing a course",
    "9.   Publishing & the course lifecycle",
    "10.  Preview & Enrollees",
    "11.  Learning Paths (bundles)",
    "12.  Skills, certificates & the marketplace",
    "13.  Recommended workflow — blank to live",
    "14.  FAQ & troubleshooting",
    "15.  Glossary",
]:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4); p.add_run(c)
page_break()

# ===========================================================================
# 1. BEFORE YOU BEGIN
# ===========================================================================
h1("1.  Before you begin", before=0)

h2("What is a course made of?")
para("A course is a small online class. It is built from a simple hierarchy — once you picture this, the editor makes complete sense:")
field_table([
    ("Course", "The whole class. Has a title, description, cover image, category, price, and a status (draft / published)."),
    ("Chapter", "A section of the course. A course has one or more chapters, in order."),
    ("Module", "A single lesson inside a chapter. This is where the actual learning content lives — text, a video, downloadable resources, and an optional reflection question."),
    ("Pause & Notice", "An optional reflective question shown after a module — no right or wrong answer, just a moment to reflect."),
    ("Integration Moment", "An optional final quiz the learner must pass to be certified."),
    ("Certificate", "Awarded automatically when the learner finishes (and passes the quiz, if there is one)."),
    ("Skills", "Badges a caregiver earns by completing the course — these show on their public profile and are searchable in the marketplace."),
], headers=("Part", "What it is"))

h2("How a learner experiences it")
steps([
    "They find the course and open its page (title, description, what's inside, price).",
    "They Enroll — instantly if free, or after paying if it's a paid course.",
    "They work through the chapters and modules in order, watching videos and reading.",
    "After some modules they may get a 'Pause & Notice' reflection.",
    "At the end, if the course has an 'Integration Moment' quiz, they must pass it.",
    "On completion they receive a downloadable certificate; caregivers also earn the course's skills.",
])

h2("Who can create courses, and where")
rich([("Only ", False), ("admins", False, True), (" create courses. Everything happens under the ", False), ("Courses", True, True),
      (" menu (web address ", False), ("/admin/courses", False, True), (").", False)])
callout("Nothing is live until you publish.", "A new course starts as a Draft — invisible to learners. You build it in private and switch it to 'Published' when ready.")
page_break()

# ===========================================================================
# 2. DASHBOARD
# ===========================================================================
h1("2.  The Courses dashboard", before=0)
rich([("Menu: ", True), ("Courses", True, True), ("   ·   ", False), ("/admin/courses", False, True)])
para("This is the home screen for all course work. It lists every course with a quick summary and the actions you'll use most.")
field_table([
    ("Course", "The course title (a star means it's Featured)."),
    ("Status", "Draft, Published, or Archived."),
    ("Chapters", "How many chapters it has."),
    ("Modules", "Total lessons across all chapters."),
    ("Enrolled", "How many learners have started it (click to see them)."),
    ("Preview", "Opens the course exactly as a learner sees it, in a new tab, without affecting anything."),
    ("Enrollees", "The list of people taking the course and their progress."),
    ("Edit", "Opens the course editor."),
], headers=("Column / link", "Meaning"))
para("Two buttons sit at the top right:")
bullet("Create course — start a brand-new course.")
bullet("Learning Paths — bundle several courses into a guided sequence (see Section 11).")
page_break()

# ===========================================================================
# 3. EDITOR AT A GLANCE
# ===========================================================================
h1("3.  The course editor at a glance", before=0)

h2("Starting a course")
steps([
    "On the Courses dashboard click 'Create course'.",
    "Give the course a Title (e.g. 'Foundations of Infant & Toddler Caregiving') and confirm.",
    "The editor opens. Your course now exists as a Draft and is saved — you can leave and come back any time.",
])

h2("The four tabs")
para("Across the top of the editor are four numbered tabs. Work through them left to right; you can jump back to any tab at any time.")
field_table([
    ("1 · Course details", "The title, description, image, classification, price, and skills."),
    ("2 · Chapters & modules", "The actual teaching content — the heart of the course."),
    ("3 · Quiz", "The optional final 'Integration Moment' quiz."),
    ("4 · Certificate", "Who signs the completion certificate and its footer."),
], headers=("Tab", "What you set there"))

h2("The save bar (always at the bottom)")
para("A bar is pinned to the bottom of the screen the whole time you edit. It contains:")
bullet("A Status drop-down — Draft, Published, or Archived.")
bullet("Back / Next buttons to move between tabs.")
bullet("A 'Save course' button.")
callout("Save often.", "Changes on any tab are only stored when you press 'Save course'. The status drop-down is also applied on save — so to publish, set it to 'Published' and then Save.")
page_break()

# ===========================================================================
# 4. TAB 1 — DETAILS
# ===========================================================================
h1("4.  Tab 1 — Course details", before=0)

h2("4.1  The details fields")
field_table([
    ("Title", "The course name."),
    ("Subtitle", "A short tagline shown under the title."),
    ("Summary", "The short blurb on the course card in listings."),
    ("Description", "The longer description on the course's own page."),
    ("Cover image URL", "A link to the cover image for the course."),
    ("Intro video", "Optional welcome video: choose the provider (None, YouTube, or Vimeo) and paste the Video URL."),
    ("Category", "What the course is about — see the list in 4.2."),
    ("Approach", "The educational philosophy — see 4.2."),
    ("Care type", "The setting it applies to: Home & Family, Small Groups, or Schools & Centers."),
    ("Age min / max (months)", "The child age range the course covers, in months (e.g. 0 to 36). Leave blank for no limit."),
    ("Est. learning (min)", "Roughly how long the whole course takes, in minutes."),
    ("Mode", "How it's delivered — free text, e.g. 'Online · Self-paced'."),
    ("Free course", "Tick for a no-cost course. Ticking hides the price fields."),
    ("Price (cents)", "Only when not free: the price in CENTS — 2500 means $25.00."),
    ("Compare-at (cents)", "Optional 'was' price in cents, shown struck-through to signal a discount."),
    ("Featured", "Tick to highlight the course to learners (shows a star)."),
    ("Allow 'skip to certification'", "Tick to let confident learners jump straight to the quiz/certificate without stepping through every module."),
], headers=("Field", "What to enter"))

h2("4.2  The exact choices for classification")
h3("Category (what it's about)")
for c in ["Health & Safety", "Healthy Habits: Hygiene, Nutrition & Sleep", "Cognitive & Motor Development",
          "Emotional and Social Development", "Language and Communication",
          "Creativity, Sensory Exploration and Expression", "Environment and Nature", "Inclusion & Diverse Needs"]:
    bullet(c)
h3("Approach (the philosophy)")
for c in ["Modern Evidence Based Approach", "Montessori", "Reggio Emilia", "Forest Schools"]:
    bullet(c)
h3("Care type (the setting)")
for c in ["Home & Family", "Small Groups", "Schools & Centers"]:
    bullet(c)
callout("Why classification matters:", "Category, Approach, Care type and the age range power the filters learners use to browse and the matching that suggests relevant courses. Fill them in even though they're optional.")

h2("4.3  The Skills panel")
para("Below the details is a 'Skills' panel — a set of clickable chips. Click the skills a caregiver earns by finishing this course. Selected chips turn coloured with a tick.")
bullet("These skills appear on a caregiver's public profile once they complete the course.")
bullet("Families and organizations search for caregivers by these skills in the marketplace, so choose accurately.")
bullet("If no chips appear, the skills library hasn't been seeded yet — ask your developer.")
page_break()

# ===========================================================================
# 5. TAB 2 — CHAPTERS & MODULES
# ===========================================================================
h1("5.  Tab 2 — Chapters & modules", before=0)
para("This is where you build the teaching content. The tab header always shows your running totals (e.g. '3 chapters · 8 modules').")

h2("5.1  Adding a chapter")
steps([
    "Click 'Add chapter' (top right of the panel).",
    "The chapter opens. Give it a Chapter title and an optional Summary.",
    "Add modules inside it (next section).",
])
para("Each chapter row shows its number, title, and module count. Use the up / down arrows to reorder chapters and the bin icon to delete one.")

h2("5.2  Adding a module")
steps([
    "Open a chapter, then click 'Add module'.",
    "The module opens for editing. Fill in the fields below.",
    "Use the up / down arrows to reorder modules within the chapter, and the bin icon to delete.",
])
para("A module can mix any of these — a module can be text only, video only, resources only, or a combination:")
field_table([
    ("Module title", "The lesson name."),
    ("Text", "The written lesson, using the rich-text editor (bold, lists, links, etc.). Optional — leave empty for a video-only or resources-only module."),
    ("Video", "Choose None, YouTube, or Vimeo."),
    ("Video URL", "The link to the video, when a provider is chosen."),
    ("Minutes", "Roughly how long the video / lesson is, in minutes."),
    ("Resources", "Downloadable or linked extras (see 5.3)."),
    ("Pause & Notice", "An optional reflection question after the module (see 5.4)."),
], headers=("Field", "What to enter"))
callout("Module tags:", "Each module row shows little tags — Text, Video, '2 resources', 'Pause & Notice' — so you can see what a module contains at a glance without opening it.")

h2("5.3  Adding resources to a module")
para("Inside the module, the 'Resources' box holds extra materials. Click 'Add' and for each resource set:")
field_table([
    ("Label", "The name the learner sees, e.g. 'Printable checklist'."),
    ("Kind", "The type: Link, YouTube, Vimeo, Google Doc, Google Drive, or File."),
    ("URL", "The web address of the resource."),
], headers=("Field", "What to enter"))
para("Use the bin icon to remove a resource.")

h2("5.4  Adding a 'Pause & Notice' reflection")
para("'Pause & Notice' is a gentle reflective question after a module — there is no right or wrong answer; it's a moment for the learner to check in with themselves.")
steps([
    "Tick the 'Pause & Notice question after this module' box inside the module.",
    "Type the question prompt (e.g. 'What feels most true for you right now?').",
    "Add answer options. For each option enter the option text and an optional Reflection / explanation shown when chosen.",
    "Tick 'highlight' on an option to gently emphasise it. Use 'Add option' for more.",
])
page_break()

# ===========================================================================
# 6. TAB 3 — QUIZ
# ===========================================================================
h1("6.  Tab 3 — Integration Moment (the quiz)", before=0)
para("The quiz is optional. A course with no quiz certifies the learner once they finish all the modules. A course with a quiz requires the learner to pass it to be certified.")

h2("6.1  How the quiz behaves")
bullet("Learners must clear the quiz to be certified.")
bullet("The pass mark defaults to 60% (you can change it).")
bullet("Learners get unlimited attempts.")
bullet("Whether each answer was right or wrong is only revealed after they pass — so it's a genuine check, not a giveaway.")

h2("6.2  Building the quiz")
steps([
    "On the Quiz tab, click 'Add quiz' (to remove it later, use 'Remove quiz').",
    "Set the Intro copy (a short line shown before the quiz) and the Pass threshold % (default 60).",
    "Click 'Add question' and type the question.",
    "Add options. For each option type the answer text and an Explanation (shown after passing).",
    "Mark the right answer by selecting its 'correct' radio button.",
    "Repeat for all questions, then Save.",
])
field_table([
    ("Intro copy", "A short introduction shown before the questions."),
    ("Pass threshold %", "The score needed to pass (default 60)."),
    ("Question", "The question prompt."),
    ("Option — correct", "The radio button marking the right answer (exactly one per question)."),
    ("Option — text", "The answer wording the learner picks from."),
    ("Option — explanation", "Why it's right/wrong, revealed only after the learner passes."),
], headers=("Field", "What to enter"))
page_break()

# ===========================================================================
# 7. TAB 4 — CERTIFICATE
# ===========================================================================
h1("7.  Tab 4 — Certificate", before=0)
para("Every learner who completes the course (and passes the quiz, if there is one) gets a downloadable certificate. This tab sets who signs it and the small print.")
field_table([
    ("Signer 1 — name", "Name of the first signatory on the certificate."),
    ("Signer 1 — title", "Their role / title, e.g. 'Program Director'."),
    ("Signer 2 — name", "Name of the second signatory (optional)."),
    ("Signer 2 — title", "Their role / title."),
    ("Footer disclaimer", "Small print shown at the bottom of the certificate."),
], headers=("Field", "What to enter"))
callout("When is it awarded?", "Automatically — there's nothing to issue by hand. The learner downloads it themselves once they've met the completion (and quiz) requirements.")
page_break()

# ===========================================================================
# 8. PRICING
# ===========================================================================
h1("8.  Pricing a course", before=0)
para("Pricing is set on Tab 1 (Course details).")
h2("Free vs paid")
bullet("Free course — tick 'Free course'. Learners enroll instantly with one click.")
bullet("Paid course — untick 'Free course' and set a Price. Learners are sent to Stripe checkout and only get access after paying.")
h2("Entering the price")
callout("Prices here are in CENTS.", "Unlike plans and events (which use dollars), course prices are entered in cents. Type 2500 for $25.00, 4999 for $49.99. The optional 'Compare-at' price (also in cents) shows a struck-through 'was' price.")
h2("What paid courses need")
bullet("Stripe must be connected (Admin → Payments). Until then, paid enrollment can't complete.")
bullet("Learners can cancel a paid course within 48 hours of buying for a full refund; access is then removed.")
page_break()

# ===========================================================================
# 9. PUBLISHING
# ===========================================================================
h1("9.  Publishing & the course lifecycle", before=0)
para("A course's visibility is controlled entirely by its Status, set in the bottom save bar.")
field_table([
    ("Draft", "Work in progress. Hidden from learners. New courses start here."),
    ("Published", "Live and visible. Learners can find, enroll in, and take the course."),
    ("Archived", "Retired. Removed from listings (existing learners' records are kept)."),
], headers=("Status", "What it means"))
h2("To publish a course")
steps([
    "Make sure the details, at least one chapter with modules, and (if used) the quiz and certificate are complete.",
    "In the bottom bar, change the Status drop-down from 'Draft' to 'Published'.",
    "Click 'Save course'. It's now live.",
])
bullet("Featured — tick 'Featured' on Tab 1 to spotlight the course to learners.")
bullet("To take a course down later, set Status back to Draft (temporary) or Archived (permanent retirement) and Save.")
page_break()

# ===========================================================================
# 10. PREVIEW & ENROLLEES
# ===========================================================================
h1("10.  Preview & Enrollees", before=0)
h2("Preview")
para("From the Courses dashboard, click 'Preview' on a course to open it exactly as a learner would see it — including videos, modules and the quiz — in a new tab. This is admin-only and does not affect your own progress or the learner data. Use it to proof-read before publishing.")
h2("Enrollees")
para("Click 'Enrollees' (or the enrolled number) to see everyone taking the course and how far they've progressed. This is your view of uptake and completion.")
page_break()

# ===========================================================================
# 11. LEARNING PATHS
# ===========================================================================
h1("11.  Learning Paths (bundles)", before=0)
rich([("Courses dashboard → ", False), ("Learning Paths", True, True), ("   ·   ", False), ("/admin/courses/bundles", False, True)])
para("A Learning Path groups several courses into a guided sequence — for example an onboarding track of three short courses taken in order.")
steps([
    "Open 'Learning Paths' from the Courses dashboard.",
    "Click to create a new path.",
    "Give it a Title, Summary and Cover image.",
    "Add the courses to include, in the order learners should take them.",
    "Optionally mark it Featured, set its Status, and Save.",
])
field_table([
    ("Title", "The name of the learning path."),
    ("Summary", "A short description of what the path covers."),
    ("Cover image", "A link to the path's cover image."),
    ("Courses", "The courses included, in sequence."),
    ("Featured", "Highlight the path to learners."),
    ("Status", "Draft (hidden) or Published (live)."),
], headers=("Field", "What to enter"))
page_break()

# ===========================================================================
# 12. SKILLS & MARKETPLACE
# ===========================================================================
h1("12.  Skills, certificates & the marketplace", before=0)
para("Courses do more than teach — for caregivers they build a verifiable, searchable professional profile.")
bullet("When a caregiver completes a course, the skills you attached on Tab 1 are added to their public profile.")
bullet("Families and organizations search and filter caregivers by these skills in the marketplace, so the skills you choose directly affect who a caregiver is matched with.")
bullet("The completion certificate is downloadable proof of the training.")
callout("Tip:", "Attach skills that genuinely reflect what the course teaches — they become claims on a caregiver's profile that families rely on.")
page_break()

# ===========================================================================
# 13. WORKFLOW
# ===========================================================================
h1("13.  Recommended workflow — blank to live", before=0)
para("A clean order of work that avoids backtracking:")
steps([
    "Create the course and give it a clear Title.",
    "Tab 1: write the Subtitle, Summary and Description; add the Cover image and (optional) intro video.",
    "Tab 1: set Category, Approach, Care type, age range and estimated time.",
    "Tab 1: decide Free or set a Price (in cents); tick Featured / skip-to-cert if wanted.",
    "Tab 1: select the Skills the course awards.",
    "Tab 2: add chapters; inside each, add modules with text and/or video, resources, and Pause & Notice questions.",
    "Tab 3: add the Integration Moment quiz if the course should be tested; mark the correct answers.",
    "Tab 4: set the certificate signers and footer.",
    "Save, then use Preview to read through it as a learner.",
    "When happy, set Status to 'Published' and Save. (For paid courses, confirm Stripe is connected first.)",
])
page_break()

# ===========================================================================
# 14. FAQ
# ===========================================================================
h1("14.  FAQ & troubleshooting", before=0)
def faq(q, a):
    p = doc.add_paragraph(); r = p.add_run(q); r.bold = True; r.font.color.rgb = PRIMARY
    p.paragraph_format.space_after = Pt(2)
    para(a, after=10)

faq("My course isn't showing to learners.",
    "Check the Status in the bottom bar is 'Published' and that you pressed 'Save course'. Draft and Archived courses are hidden.")
faq("I set a price but nobody can buy it.",
    "Paid courses need Stripe connected (Admin → Payments). Until Stripe is set up, only free courses can be enrolled in.")
faq("The price looks wrong — it's 100x too big.",
    "Course prices are in cents. Enter 2500 for $25.00, not 25.")
faq("There are no skill chips to choose.",
    "The skills library hasn't been seeded yet. Ask your developer to add it; everything else still works.")
faq("Do I have to add a quiz?",
    "No. Without a quiz, learners are certified once they complete all modules. Add a quiz only when you want a pass/fail check.")
faq("Can I see the course as a student before publishing?",
    "Yes — use 'Preview' on the Courses dashboard. It opens the real learner view in a new tab without affecting any data.")
faq("How do I reorder chapters or modules?",
    "Use the up / down arrows on each chapter or module row. The order you see is the order learners get.")
faq("What happens to a course I archive?",
    "It disappears from listings and can't be newly enrolled in, but existing learners' progress and certificates are preserved.")
page_break()

# ===========================================================================
# 15. GLOSSARY
# ===========================================================================
h1("15.  Glossary", before=0)
field_table([
    ("Chapter", "A section of a course that groups related modules."),
    ("Module", "A single lesson: text and/or video, plus optional resources and a reflection."),
    ("Pause & Notice", "A reflective question after a module — no right or wrong answer."),
    ("Integration Moment", "The course's final quiz; learners must pass it to be certified."),
    ("Pass threshold", "The score needed to pass the quiz (default 60%)."),
    ("Certificate", "The downloadable proof of completion, awarded automatically."),
    ("Skill", "A badge a caregiver earns from a course, searchable in the marketplace."),
    ("Learning Path", "A bundle of courses taken as a guided sequence."),
    ("Category / Approach / Care type", "The classification that powers browsing filters and matching."),
    ("Compare-at price", "An optional 'was' price shown struck-through to signal a discount."),
    ("Draft / Published / Archived", "Hidden work-in-progress / live and visible / retired."),
    ("Featured", "A spotlight flag that highlights a course or path to learners."),
    ("Preview", "An admin-only view of a course exactly as a learner sees it."),
    ("Cents", "The unit course prices are entered in — 2500 = $25.00."),
], headers=("Term", "Meaning"))

doc.add_paragraph().paragraph_format.space_before = Pt(12)
end = doc.add_paragraph(); end.alignment = WD_ALIGN_PARAGRAPH.CENTER
er = end.add_run("— End of guide —"); er.italic = True; er.font.color.rgb = MUTED

# ===========================================================================
root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
out = os.path.join(root, "The-Raising-Club-Course-Creation-Guide.docx")
doc.save(out)
print("Saved:", out)
