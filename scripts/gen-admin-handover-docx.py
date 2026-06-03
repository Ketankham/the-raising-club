# Generates the Word (.docx) Admin Handover Guide.
# Run:  python scripts/gen-admin-handover-docx.py
# Output: ../The-Raising-Club-Admin-Handover.docx  (project root, not in the git repo)

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK     = RGBColor(0x2B, 0x2B, 0x2B)
PRIMARY = RGBColor(0xC0, 0x53, 0x3B)
GREEN   = RGBColor(0x5E, 0x7B, 0x3A)
MUTED   = RGBColor(0x6B, 0x6B, 0x6B)
AMBER   = RGBColor(0x9A, 0x6B, 0x00)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def h1(text, before=18):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = PRIMARY
        r.font.size = Pt(18)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(6)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = INK
        r.font.size = Pt(13.5)
    p.paragraph_format.space_before = Pt(12)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.color.rgb = GREEN
        r.font.size = Pt(11.5)
    p.paragraph_format.space_before = Pt(8)
    return p

def para(text="", bold=False, italic=False, color=None, size=None, after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color is not None: r.font.color.rgb = color
    if size is not None:  r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(after)
    return p

def rich(parts, after=6):
    """parts = list of (text, bold, italic) tuples on one paragraph."""
    p = doc.add_paragraph()
    for t in parts:
        text = t[0]; bold = t[1] if len(t) > 1 else False; italic = t[2] if len(t) > 2 else False
        r = p.add_run(text); r.bold = bold; r.italic = italic
    p.paragraph_format.space_after = Pt(after)
    return p

def bullet(text, sub=False):
    p = doc.add_paragraph(style="List Bullet" + (" 2" if sub else ""))
    p.add_run(text)
    p.paragraph_format.space_after = Pt(2)
    return p

def steps(items):
    for i, it in enumerate(items, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}.  ")
        run.bold = True
        run.font.color.rgb = PRIMARY
        p.add_run(it)
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.space_after = Pt(3)

def callout(label, text, color=AMBER):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"{label}  ")
    r.bold = True; r.font.color.rgb = color
    r2 = p.add_run(text); r2.italic = True; r2.font.color.rgb = MUTED

def field_table(rows, headers=("Field", "What it is / what to enter")):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        rp = hdr[i].paragraphs[0].add_run(htext)
        rp.bold = True; rp.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); rp.font.size = Pt(10)
        _shade(hdr[i], "C0533B")
    for name, desc in rows:
        cells = t.add_row().cells
        cells[0].text = ""
        r0 = cells[0].paragraphs[0].add_run(name); r0.bold = True; r0.font.size = Pt(10)
        cells[1].text = ""
        cells[1].paragraphs[0].add_run(desc).font.size = Pt(10)
    # column widths
    for row in t.rows:
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(4.6)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def page_break():
    doc.add_page_break()

# ===========================================================================
# COVER
# ===========================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("The Raising Club")
tr.font.size = Pt(30); tr.bold = True; tr.font.color.rgb = PRIMARY
title.paragraph_format.space_before = Pt(120)

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Admin Handover Guide"); sr.font.size = Pt(20); sr.font.color.rgb = INK

sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
s2 = sub2.add_run("A complete, step-by-step manual for managing users, events, courses, plans and payments")
s2.italic = True; s2.font.color.rgb = MUTED; s2.font.size = Pt(11)

meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(40)
mr = meta.add_run("For new administrators · No technical background required")
mr.font.color.rgb = MUTED; mr.font.size = Pt(10)
page_break()

# ===========================================================================
# CONTENTS
# ===========================================================================
h1("What's in this guide", before=0)
contents = [
    "1.  Before you begin — what 'admin' is, signing in, the menu",
    "2.  User management — invite, search, manage, deactivate, assign plans",
    "3.  Plan management — create and edit membership plans",
    "4.  Payment settings — connecting Stripe (test & live)",
    "5.  Event management — create, publish, roster, paid events",
    "6.  Course management — build courses, pricing, publishing",
    "7.  Notifications — the message templates users receive",
    "8.  How families & memberships fit together",
    "9.  Frequently asked questions & troubleshooting",
    "10. Glossary of terms",
]
for c in contents:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    p.add_run(c)
page_break()

# ===========================================================================
# 1. BEFORE YOU BEGIN
# ===========================================================================
h1("1.  Before you begin", before=0)

h2("What does 'admin' mean here?")
para("The Raising Club is used by four kinds of people. Every account has exactly one of these roles:")
field_table([
    ("Parent / Family", "A parent or family member. Can join events, take courses, connect with other families, and post care jobs."),
    ("Caregiver / Educator", "A childcare professional. Builds a public profile, applies to jobs, takes training."),
    ("Organization", "A centre or program (e.g. a daycare). Posts jobs, hosts events, manages staff."),
    ("Admin", "You. Full control of the platform — every user, event, course, plan and payment setting."),
], headers=("Role", "Who they are"))
para("As an admin you can see and manage everything. The other three roles only ever see their own data.")

h2("Signing in")
steps([
    "Go to the website and click 'Sign in' (top right).",
    "Enter the email and password of your admin account.",
    "After signing in you will land on the admin area automatically.",
])
callout("You need an admin account first.",
        "Admin accounts are created by your developer (a one-time setup step). If you don't have one yet, ask them to create it for you, or have an existing admin invite you with the role 'Admin'.")

h2("The admin menu")
para("Once signed in as an admin, a menu runs down the left-hand side of the screen. These are the only six places you need:")
field_table([
    ("Users", "Everyone who has an account. Invite, search, manage, and assign plans."),
    ("Events", "Create and run events (workshops, classes, play sessions). Shared with organization owners."),
    ("Plans", "The membership plans shown on the pricing page. Create, edit, price, and see who is subscribed."),
    ("Courses", "Training courses and learning paths. Build content, set pricing, publish."),
    ("Notifications", "The wording of the in-app messages users receive (e.g. 'You're registered')."),
    ("Payments", "Where you connect Stripe so the platform can take card payments."),
], headers=("Menu item", "What it controls"))

h2("Three things that are true everywhere")
bullet("Saving: changes are never live until you press a Save / Create button. Nothing auto-saves.")
bullet("Invitations are links, not emails. Email sending is not switched on yet, so when you invite someone the system gives you a link to copy and send yourself (by text, WhatsApp, email, etc.).")
bullet("Test mode vs Live mode: payments can run in 'Test' mode (fake cards, safe to practise) or 'Live' mode (real money). You choose which in Payments. Always practise in Test first.")
page_break()

# ===========================================================================
# 2. USER MANAGEMENT
# ===========================================================================
h1("2.  User management", before=0)
rich([("Menu: ", True), ("Users", True, True), ("   ·   Web address ends in ", False), ("/admin", False, True)])

h2("2.1  The overview at the top")
para("The coloured boxes across the top are a live count of your community:")
field_table([
    ("Total users", "Everyone who has ever started an account."),
    ("Active", "Accounts that are not deactivated."),
    ("Onboarded", "People who finished the sign-up questionnaire."),
    ("Parents / Caregivers / Organizations", "How many of each role you have."),
], headers=("Box", "Meaning"))

h2("2.2  Inviting a new user")
steps([
    "In the 'Invite a user' box, type the person's email address.",
    "Choose their role from the drop-down (Parent, Caregiver, Organization, or Admin).",
    "Click 'Send invite'.",
    "The system creates an invitation and copies a private link to your clipboard (you'll see 'link copied' at the bottom).",
    "Paste that link to the person however you like. When they open it, they create their account in that role.",
])
callout("Tip:", "Pending invitations appear just below the invite box. You can 'Copy link' again at any time, or 'Revoke' an invitation you no longer want to honour.")

h2("2.3  Finding a user")
para("The 'All users' table lists everyone, newest first. Use the search box (top right of the table) to filter by name, email, or role instantly.")
para("Each row shows:")
field_table([
    ("User", "Name and email."),
    ("Role", "Parent, Caregiver, Organization, or Admin."),
    ("Onboarding", "'Complete', 'In progress', or '—' (not started)."),
    ("Email", "'Confirmed' once they verified their email address, otherwise 'Unconfirmed'."),
    ("Registered", "The date they created the account."),
    ("Status", "'Active' (green) or 'Deactivated' (red)."),
    ("Actions", "'Manage' opens their full profile; 'Deactivate' / 'Reactivate' toggles access."),
], headers=("Column", "Meaning"))

h2("2.4  Deactivating and reactivating")
para("Deactivating is the safe alternative to deleting. The account and all its data stay, but the person can no longer sign in or be seen by others.")
steps([
    "Find the user in the table.",
    "Click 'Deactivate' on their row (admins cannot be deactivated this way).",
    "To restore them later, click 'Reactivate' on the same row.",
])

h2("2.5  Managing one user in detail")
rich([("Click ", False), ("Manage", False, True), (" on any user to open their detail page (", False), ("/admin/users/<id>", False, True), ("). It has four areas:", False)])

h3("Account (read-only)")
para("Reference facts you cannot edit here: account type, email, email status, onboarding status, registration date, and active/deactivated status.")

h3("Personal details (editable)")
para("Edit the person's First name, Last name, Preferred name, Phone, and ZIP / postcode. Click 'Save changes'.")

h3("Membership (quick plan record)")
para("A simple way to record which plan a user has selected, with a Monthly / Annual toggle. This only records the choice — it does not take payment. For real billing or free comp access, use the next section.")

h3("Plans & entitlement (the important one)")
para("This panel shows the user's live access and lets you grant a plan for free (a 'comp grant'). It has three parts:")
bullet("Snapshot — a coloured badge with their current status (active, trialing, comp, past_due, canceled, or none), the plan name, and the date it runs until.")
bullet("Assignments — every plan ever attached to this person (whether bought through Stripe or granted by you), with status and date window. Comp grants you made can be 'Revoked' here.")
bullet("Comp grant — give this person (or their family / organization) a plan at no charge.")
para("To give someone a free plan:", bold=True)
steps([
    "Choose the Plan from the drop-down (only plans for that person's role appear).",
    "Pick a Billing label — Monthly or Annual (this is only a label; nothing is charged).",
    "Set the Start date.",
    "Set the End date (required — use the +1m / +1y shortcuts for one month or one year).",
    "Optionally add a Note (e.g. 'Beta tester comp').",
    "Click 'Assign comp plan'. Their access turns on immediately and lasts until the End date.",
])
callout("Good to know:", "For a family (parent) the plan is attached to their whole household, so every adult they've invited gets the access too. Assigning a new comp plan automatically replaces any existing active one for that person.")
callout("Important:", "Comp grants never involve Stripe and never charge a card. They are purely a gift of access for the window you set.")
page_break()

# ===========================================================================
# 3. PLAN MANAGEMENT
# ===========================================================================
h1("3.  Plan management", before=0)
rich([("Menu: ", True), ("Plans", True, True), ("   ·   Web address ends in ", False), ("/admin/plans", False, True)])

para("Plans are the membership tiers shown on the public pricing page and in each user's Settings. Everything about them — names, prices, features, who they're for — is edited here. Changes appear on the website immediately after saving.")

h2("3.1  How plans are organised")
para("Every plan belongs to one of three audiences, matching the three tabs on the pricing page:")
field_table([
    ("Caregiver & Educator", "Plans for childcare professionals."),
    ("Family", "Plans for parents and families. These carry 'adult seats' — how many adults can share one family membership."),
    ("Centers & Programs", "Plans for organizations. These carry 'staff seats'."),
], headers=("Audience", "Who it's for"))

h2("3.2  Viewing and hiding plans")
para("The Plans page groups every plan under its audience. On each plan you can:")
bullet("Click the name to edit it.")
bullet("Tick / untick 'Active' to instantly show or hide it on the public pricing page (hidden plans show a grey 'Hidden' tag).")
bullet("Click 'Subscribers' to see who is currently on that plan.")

h2("3.3  Creating or editing a plan")
steps([
    "Click 'New plan' (top right), or click an existing plan's name to edit it.",
    "Fill in the fields below.",
    "Click 'Create plan' / 'Save changes'.",
])
para("Full field reference:", bold=True)
field_table([
    ("Stable key", "A short code, e.g. family_access. Set once when creating; it cannot be changed later because it is stored against every subscription. Lowercase letters, numbers and underscores only."),
    ("Audience", "Caregiver & Educator, Family, or Centers & Programs (see 3.1)."),
    ("Name", "The plan's display name, e.g. 'Family Access'."),
    ("Badge", "Optional ribbon on the card, e.g. 'Most popular' or 'Best value'. Leave blank for none."),
    ("Subtitle", "A short tagline under the name, e.g. 'Go deeper'."),
    ("Description", "A sentence or two describing who the plan is for."),
    ("CTA label", "The button text on the pricing card, e.g. 'Get Family Access'."),
    ("Position", "Sort order within its audience (0 is first)."),
    ("Highlighted card", "Tick to make this card stand out as the recommended choice."),
    ("Free plan", "Tick if this is a no-cost starter plan. Hides the price fields."),
    ("Custom pricing", "Tick for 'contact us' plans with no fixed price (e.g. enterprise)."),
    ("Active", "Tick to show the plan publicly; untick to hide it."),
    ("Custom label", "Only for custom-pricing plans: the text shown instead of a price."),
    ("Monthly ($)", "Monthly price in dollars (e.g. 29). Leave blank for free / custom plans."),
    ("Annual ($/yr)", "The full yearly price in dollars (e.g. 296). Shown when a visitor toggles to Annual."),
    ("Unit", "Optional suffix shown after the price, e.g. 'per site' for organization plans."),
    ("Adult seats", "Family plans only: how many adults can share this membership (e.g. Essentials 1, Access 3, Club+ 5)."),
    ("Staff seats", "Organization plans only: how many staff the plan covers."),
    ("Features", "The bullet list on the card. Each feature has a bold Label and a longer Body. Use 'Add feature' for more, or the bin icon to remove one."),
    ("Stripe price IDs", "Found under 'Stripe price IDs'. The Product ID, Monthly price ID and Annual price ID copied from your Stripe dashboard — see the box below."),
], headers=("Field", "What to enter"))
callout("Paid plans need Stripe price IDs.",
        "A plan cannot be purchased until you paste its Stripe price ID for the chosen billing period. Create the product & prices in Stripe, then copy the price IDs (they start with 'price_') into the plan here. Free plans and comp grants do not need this.")

h2("3.4  Seeing who is subscribed")
para("From the Plans list, click 'Subscribers' on a plan to open its subscriber list. Each row shows the person / family / organization, their status, whether it's a Stripe subscription or a comp grant, the billing period, and the start and renewal/end dates.")
page_break()

# ===========================================================================
# 4. PAYMENT SETTINGS
# ===========================================================================
h1("4.  Payment settings (Stripe)", before=0)
rich([("Menu: ", True), ("Payments", True, True), ("   ·   Web address ends in ", False), ("/admin/settings/payments", False, True)])

para("Stripe is the service that takes card payments. The platform comes with payments switched off; this page is where you connect your Stripe account. You enter the keys once and can update them any time — no developer or redeploy needed.")

h2("4.1  Test mode vs Live mode")
para("At the top you choose the Active mode:")
field_table([
    ("Test", "Uses Stripe's test keys and fake test cards. No real money moves. Use this to practise and check everything works."),
    ("Live", "Uses your real Stripe keys. Real cards are charged. A warning is shown to remind you."),
], headers=("Mode", "What it does"))
para("You can save both sets of keys and flip between them with the toggle. The website always uses whichever mode is selected.")

h2("4.2  Entering your keys")
para("You'll find these three values in your Stripe dashboard, under Developers. There is a Test set and a Live set; enter each in its own section here.")
field_table([
    ("Publishable key", "Starts with 'pk_'. Safe to share; used by the checkout page."),
    ("Secret key", "Starts with 'sk_'. Private. Shown as dots once saved — leave blank to keep the existing one, or type a new value to replace it."),
    ("Webhook signing secret", "Starts with 'whsec_'. Lets the platform trust messages from Stripe (see 4.3)."),
], headers=("Field", "What to enter"))
steps([
    "Choose the section you're filling in (Test keys or Live keys).",
    "Paste the Publishable key, Secret key, and Webhook signing secret.",
    "Click 'Save keys'. Saved secrets show as dots with the last 4 characters.",
    "Click 'Test … connection' to confirm the keys are valid — you'll see a green success message.",
])

h2("4.3  Setting up the webhook (one-time)")
para("A webhook lets Stripe tell the platform when a payment succeeds, a subscription renews, or a card fails. Without it, purchases won't be recorded.")
steps([
    "On this page, copy the 'Webhook endpoint' URL shown in the grey box.",
    "In your Stripe dashboard, go to Developers → Webhooks → 'Add endpoint' and paste that URL.",
    "Subscribe the endpoint to subscription, invoice, and checkout events (your developer can confirm the exact list).",
    "Stripe will show a 'Signing secret' (starts with whsec_). Copy it.",
    "Paste it into the matching 'Webhook signing secret' field here and click 'Save keys'.",
])

h2("4.4  Going live — checklist")
bullet("Enter and save your Live publishable, secret and webhook keys.")
bullet("Add the Live webhook endpoint in Stripe and paste its signing secret here.")
bullet("In Plans, add the Stripe price IDs to each paid plan.")
bullet("Switch the Active mode to 'Live'.")
bullet("Make one small real purchase to confirm money arrives in Stripe.")
callout("Note:", "There is also one technical setting (a 'service role key') your developer must add to the hosting environment for webhooks to work. If purchases aren't being recorded, that's the first thing to check with them.")
page_break()

# ===========================================================================
# 5. EVENT MANAGEMENT
# ===========================================================================
h1("5.  Event management", before=0)
rich([("Menu: ", True), ("Events", True, True), ("   ·   Web address ends in ", False), ("/manage/events", False, True)])

para("Events are workshops, classes, play sessions and series that families and caregivers can register for. As an admin you can manage every event; organization owners can manage their own.")

h2("5.1  The events list")
para("The Events page lists your events. From each you can open its 'Roster' (who's coming), 'Messages' (talk to attendees), and 'Edit'. Use 'Create event' to make a new one.")

h2("5.2  Creating an event")
steps([
    "Click 'Create event'.",
    "Fill in the fields (reference below). Required fields are marked with a red asterisk.",
    "Set Status to 'Draft' while you work, then 'Published' when ready for people to see and register.",
    "Save.",
])
para("Field reference:", bold=True)
field_table([
    ("Title", "The event name."),
    ("Short summary", "One line shown on event cards and lists."),
    ("What to expect", "The full description on the event page."),
    ("Hero image URL", "Link to a banner image for the event."),
    ("Hosting organization", "If you manage more than one organization, which one is hosting (admins can host on behalf of any)."),
    ("Event style", "Guided class, Open play, Workshop, or Ongoing series."),
    ("Who attends", "Children with an adult, Children drop-off (no adult stays), or Adults only."),
    ("Min / Max age (months)", "The suitable age range for children, in months. Leave blank for no limit."),
    ("Starts / Ends", "Date and time the event runs."),
    ("Time zone", "The event's local time zone."),
    ("How you'll join / Location", "For in-person: Address plus arrival / parking notes. For online: the Platform (e.g. Zoom) and a Join URL."),
    ("Host details", "The person leading it — Name, Role / title, Avatar image URL, and a short Bio."),
    ("Resources", "Optional links (Label + URL) shared with attendees."),
    ("Pricing", "Included (free), Paid, or Donation (pay-what-you-can). See 5.3."),
    ("Price (USD)", "The amount, when Pricing is Paid."),
    ("Child capacity", "Maximum number of children. Blank means unlimited. Capacity is counted per child, not per family."),
    ("Adult capacity", "Maximum number of adults. Blank means unlimited."),
    ("Requires approval", "Tick if you want to approve each registration before it's confirmed."),
    ("Visibility", "Public (anyone can find it) or Private (only people with the link)."),
    ("Status", "Draft, Published, Full, Cancelled, Completed, or Archived."),
], headers=("Field", "What to enter"))

h2("5.3  Pricing models explained")
field_table([
    ("Included", "Free. Registration is instant (or pending if you require approval). No payment step."),
    ("Paid", "Costs money. Attendees are sent to Stripe checkout and are only confirmed after paying."),
    ("Donation", "Pay-what-you-can. Treated like a paid event with a suggested amount."),
], headers=("Pricing", "Behaviour"))
callout("Paid events need Stripe.", "Just like plans, paid and donation events only work once Stripe is connected (Section 4). Until then, free events work normally.")

h2("5.4  Running the event")
bullet("Roster — see and manage who's registered. If you ticked 'Requires approval', you approve people here. You can also check children in on the day.")
bullet("Messages — send updates to everyone registered.")
bullet("Cancelling — set Status to 'Cancelled'. For paid events, refunds and credits are handled through Stripe / the registration.")
page_break()

# ===========================================================================
# 6. COURSE MANAGEMENT
# ===========================================================================
h1("6.  Course management", before=0)
rich([("Menu: ", True), ("Courses", True, True), ("   ·   Web address ends in ", False), ("/admin/courses", False, True)])

para("Courses are the training content. The list shows each course's status, number of chapters and modules, and how many people are enrolled. From each row you can Preview it as a learner, see Enrollees, or Edit.")

h2("6.1  Creating a course")
para("Click 'Create course'. The editor has four tabs along the top — work through them left to right.")

h3("Tab 1 — Course details")
field_table([
    ("Title / Subtitle", "The course name and a short secondary line."),
    ("Summary / Description", "A short blurb and the full description."),
    ("Cover image URL", "Link to the course's cover image."),
    ("Intro video", "Optional welcome video — choose the provider and paste the Video URL."),
    ("Category / Approach / Care type", "Classification used for filtering and matching."),
    ("Age min / max (months)", "The child age range the course is about."),
    ("Est. learning (min)", "Roughly how long the course takes, in minutes."),
    ("Mode", "How it's delivered, e.g. 'Online · Self-paced'."),
    ("Price (cents)", "The price in CENTS (e.g. 2500 = $25). Set 0 and mark it free for a free course."),
    ("Compare-at (cents)", "Optional 'was' price, in cents, shown struck-through."),
], headers=("Field", "What to enter"))
callout("Prices are in cents here.", "Unlike plans and events (which use dollars), the course price is entered in cents. 2500 means $25.00.")

h3("Tab 2 — Chapters & modules")
para("This is the course content itself.")
steps([
    "Click 'Add chapter' and give it a Title (and optional Summary).",
    "Inside a chapter, click to add a Module. Each module has a Title and is either Text or Video.",
    "For video modules, paste the Video URL and optional Minutes; for text, write the body.",
    "Optionally add Resources (downloadable links) and a reflection question to a module.",
    "Use the Move up / Move down and Delete controls to reorder or remove items.",
])

h3("Tab 3 — Quiz")
para("Optionally add quiz questions learners answer to test understanding and earn their certificate.")

h3("Tab 4 — Certificate")
para("Configure the certificate awarded on completion: Intro copy, Signer 1 and Signer 2 (name + title), and a Footer disclaimer.")

h2("6.2  Pricing, publishing and featuring")
bullet("Free vs paid — a course with price 0 (or marked free) enrolls instantly. A paid course sends the learner to Stripe checkout and only grants access after payment.")
bullet("Status — Draft (hidden), Published (live), or Archived (retired).")
bullet("Featured — a starred course is highlighted to learners.")
bullet("Refunds — learners can cancel a paid course within 48 hours of buying for a full refund.")

h2("6.3  Learning Paths")
para("From the Courses page, 'Learning Paths' lets you bundle several courses together into a guided sequence.")

h2("6.4  Checking on learners")
para("Click a course's 'Enrollees' to see who is taking it, and 'Preview' to experience the course exactly as a learner would (without affecting your own progress).")
page_break()

# ===========================================================================
# 7. NOTIFICATIONS
# ===========================================================================
h1("7.  Notifications", before=0)
rich([("Menu: ", True), ("Notifications", True, True), ("   ·   Web address ends in ", False), ("/admin/notifications", False, True)])

para("This is where the wording of the messages users receive is managed — for example 'You're registered', 'Payment received', or 'Your membership is active'. They appear in the bell menu inside the app.")
bullet("Messages are grouped by category: Courses, Events, Marketplace, and General.")
bullet("Open any message type to edit its title and body, and to toggle whether it's shown in-app.")
bullet("Use the token chips (e.g. {{eventName}}) to drop in details that get filled automatically — there's a live preview with sample values.")
callout("Email is parked.", "The in-app messages work today. Email delivery of these messages is switched off for now, so users see them inside the app only.")
page_break()

# ===========================================================================
# 8. FAMILIES & MEMBERSHIPS
# ===========================================================================
h1("8.  How families & memberships fit together", before=0)
para("This background helps you answer member questions confidently.")

h2("A family is one shared membership")
bullet("A parent has a 'household' (their Raising Club). One family plan covers the whole household.")
bullet("The plan's 'adult seats' decide how many adults the family can invite (e.g. 3 adults on Family Access).")
bullet("The parent invites co-parents, grandparents, etc. by link from their own Settings → Your Raising Club. Invited adults share the family's membership and can post jobs and message other families.")
bullet("If a family tries to invite more adults than their plan allows, they're prompted to upgrade.")

h2("Where a member's access comes from")
bullet("A bought subscription (through Stripe), OR")
bullet("A comp grant you gave them (Section 2.5), OR")
bullet("For invited adults, the access they inherit from their family's plan.")
para("In every case the user's current access and its end date show on their detail page (Plans & entitlement) and in their own Settings.")
page_break()

# ===========================================================================
# 9. FAQ
# ===========================================================================
h1("9.  Frequently asked questions", before=0)

def faq(q, a):
    p = doc.add_paragraph(); r = p.add_run(q); r.bold = True; r.font.color.rgb = PRIMARY
    p.paragraph_format.space_after = Pt(2)
    para(a, after=10)

faq("A user can't pay / checkout doesn't work.",
    "Check three things: (1) Stripe keys are saved in Payments and 'Test connection' passes; (2) the Active mode is the one you intend; (3) for a plan, its Stripe price ID is filled in. If still failing, ask your developer to confirm the webhook and service-role key are set.")
faq("I invited someone but they got no email.",
    "That's expected — email is not switched on. Inviting copies a link to your clipboard; you send that link to the person yourself.")
faq("What's the difference between a comp grant and a subscription?",
    "A subscription is paid through Stripe and renews automatically. A comp grant is free access you give for a fixed window — no card, no charge. Both appear on the user's Plans & entitlement panel.")
faq("How do I stop someone using the platform without deleting them?",
    "Deactivate them (Users table or their detail page). They can't sign in or be seen, but their data is kept and you can reactivate any time.")
faq("I changed a plan's price — when does it show?",
    "Immediately after saving. The public pricing page and Settings read plans live from the database.")
faq("A plan won't appear on the pricing page.",
    "Make sure 'Active' is ticked on the plan. Hidden (inactive) plans never show publicly.")
faq("How do I give a whole family free access?",
    "Open the parent's detail page and use the Comp grant — it attaches to their household, so every invited adult gets it too.")
faq("Can I practise without charging real cards?",
    "Yes. Keep Payments in 'Test' mode and use Stripe's test card numbers. Switch to 'Live' only when you're ready to take real money.")
page_break()

# ===========================================================================
# 10. GLOSSARY
# ===========================================================================
h1("10.  Glossary", before=0)
field_table([
    ("Onboarding", "The sign-up questionnaire a new user completes to set up their profile."),
    ("Role", "What kind of account someone has: Parent, Caregiver, Organization, or Admin."),
    ("Plan", "A membership tier (free or paid) shown on the pricing page."),
    ("Subscription", "A paid plan billed automatically through Stripe."),
    ("Comp grant", "Free plan access an admin gives a user for a set period — no payment."),
    ("Entitlement", "What a user is currently allowed to access, based on their plan/comp/family."),
    ("Household", "A family group that shares one family membership."),
    ("Seat", "One slot in a family or organization plan — how many adults/staff it covers."),
    ("Stripe", "The payment provider that processes cards and subscriptions."),
    ("Test / Live mode", "Practice payments (fake) vs real payments (real money)."),
    ("Webhook", "An automatic message from Stripe telling the platform a payment happened."),
    ("Deactivate", "Switch off an account without deleting it."),
    ("Roster", "The list of people registered for an event."),
    ("Draft / Published", "Hidden work-in-progress vs live and visible to users."),
], headers=("Term", "Meaning"))

doc.add_paragraph().paragraph_format.space_before = Pt(12)
end = doc.add_paragraph(); end.alignment = WD_ALIGN_PARAGRAPH.CENTER
er = end.add_run("— End of guide —"); er.italic = True; er.font.color.rgb = MUTED

# ===========================================================================
out = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                   "The-Raising-Club-Admin-Handover.docx")
doc.save(out)
print("Saved:", out)
