# Generates the Word (.docx) handover doc for Course Creation.
# Run:  python scripts/gen-course-handover-docx.py
# Output: ../Course-Creation-Handover.docx  (project root, not in the git repo)

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

INK = RGBColor(0x2B, 0x2B, 0x2B)
PRIMARY = RGBColor(0xC0, 0x53, 0x3B)
MUTED = RGBColor(0x6B, 0x6B, 0x6B)

doc = Document()

# Base styles
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK

def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = PRIMARY
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = INK
    return p

def para(text="", italic=False, bold=False, color=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    if color: r.font.color.rgb = color
    if size: r.font.size = Pt(size)
    return p

def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def numbered(text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

# ---------------------------------------------------------------------------
# TITLE
# ---------------------------------------------------------------------------
title = doc.add_heading("The Raising Club — Course Creation Handover", level=0)
for r in title.runs:
    r.font.color.rgb = PRIMARY
para("A guide to creating courses as an admin: every screen, every term, the learner "
     "experience, the quiz, and certification. Includes the new admin-only course Preview.",
     italic=True, color=MUTED)
para("Audience: platform admins and the product team. Environment: the admin console at /admin/courses.",
     italic=True, color=MUTED, size=9)

# ---------------------------------------------------------------------------
h1("1. Where course creation lives & who can do it")
bullet(" Courses are created and managed by platform admins only. Everything is under the admin console at /admin/courses (guarded so non-admins are redirected away).", bold_lead="Admin-only.")
bullet(" Backend is Supabase. A course is a tree: Course → Chapters → Modules (each module may have resources and a “Pause & Notice” question) → an optional final Quiz → a Certificate configuration.", bold_lead="Data model.")
bullet(" Courses can also be grouped into Learning Paths (bundles) at /admin/courses/bundles.", bold_lead="Learning Paths.")

h1("2. How to get there (navigation)")
h2("Admin (authoring)")
table(
    ["Action", "Where to click", "URL"],
    [
        ["Open the course list", "Admin sidebar → Courses", "/admin/courses"],
        ["Create a course", "Courses page → “Create course” (top right)", "/admin/courses/new"],
        ["Edit a course", "Courses table → “Edit” on a row", "/admin/courses/[id]/edit"],
        ["Preview a course (admin-only)", "Courses table → “Preview”, or “Preview” in the editor header", "/admin/courses/[id]/preview"],
        ["Learning Paths (bundles)", "Courses page → “Learning Paths”", "/admin/courses/bundles"],
    ],
    widths=[2.2, 3.2, 2.2],
)
h2("Learner (the published experience)")
table(
    ["Screen", "What it is", "URL"],
    [
        ["Course catalog", "Public list of published courses", "/courses"],
        ["Course detail / enroll", "Intro page; enroll to start", "/courses/[slug]"],
        ["Course player", "Lessons + “Pause & Notice” (after enrolling)", "/courses/[slug]"],
        ["Integration Moment", "The final quiz", "/courses/[slug]/quiz"],
        ["Certificate", "Earned certificate", "/courses/[slug]/certificate"],
        ["My courses", "A learner’s enrolled courses", "/courses/my  (and /dashboard/courses)"],
        ["Certificate verify", "Public verification by token", "/verify/[token]"],
    ],
    widths=[1.9, 3.1, 2.6],
)

# ---------------------------------------------------------------------------
h1("3. Creating a course — step by step")
para("Creating a course is a two-stage flow: name it, then build it in the tabbed editor.", color=MUTED, italic=True)

h2("Stage 1 — Name the course")
numbered(" Go to /admin/courses and click “Create course”.", )
numbered(" Enter a Course title and click “Create & continue”. This creates a draft and opens the editor.")

h2("Stage 2 — The editor (four tabs)")
para("The editor has four numbered tabs. Use “Save course” (bottom bar) often — it saves the whole course tree. "
     "The status selector (Draft / Published / Archived) is also in the bottom bar.", color=MUTED, italic=True)

para("Tab 1 — Course details", bold=True, color=PRIMARY)
bullet("Title, Subtitle, Summary (card blurb), Description (detail page).")
bullet("Cover image URL; Intro video (provider + URL).")
bullet("Category, Approach, Care type (taxonomy — drives filtering/positioning).")
bullet("Age min / Age max (months), Est. learning (minutes), Mode (e.g. “Online · Self-paced”).")
bullet("Pricing: Free toggle, Price (cents), Compare-at (cents). (Payment is schema-level; courses are typically free today.)")

para("Tab 2 — Chapters & modules (the curriculum)", bold=True, color=PRIMARY)
bullet("Add chapter → give it a title (and optional summary). Chapters are the sidebar groups in the player.")
bullet("Inside a chapter, Add module. A module can have: a title, rich-text body, a video (YouTube/Vimeo) + minutes, file/link Resources, and optionally a “Pause & Notice” reflection question.")
bullet("Order chapters and modules with the up/down arrows.")

para("Tab 3 — Quiz (the “Integration Moment”)", bold=True, color=PRIMARY)
bullet("Click “Add quiz”, set Intro copy and Pass threshold % (default 60).")
bullet("Add question → write the prompt, add options, and mark the correct one. Add an Explanation per option (shown to learners after they pass).")
bullet("Courses without a quiz certify on completing all modules; courses with a quiz require passing it.")

para("Tab 4 — Certificate", bold=True, color=PRIMARY)
bullet("Configure Signer 1 / Signer 2 (name + title) and a Footer disclaimer that appear on the issued certificate.")

para("Then: set status to Published (bottom bar) and Save. Publishing makes the course visible on /courses. "
     "Tip: use the new Preview (next section) before publishing.", italic=True)

# ---------------------------------------------------------------------------
h1("4. Glossary — every term, what it means, what it affects")

h2("Course-level")
table(
    ["Term", "Meaning / effect"],
    [
        ["Title / Subtitle", "Course name and short tagline shown on the card and detail page."],
        ["Summary", "The blurb on the course card in the catalog."],
        ["Description", "The longer description on the course detail page."],
        ["Cover image URL", "Image shown on the card and detail page."],
        ["Intro video", "Optional video on the detail page (and fallback in the player)."],
        ["Category / Approach", "Taxonomy used to organize and position the course."],
        ["Care type", "Home & Family / Small Groups / Schools & Centers — shared with the marketplace care types."],
        ["Age min / max (months)", "Intended child age range for the course."],
        ["Est. learning (minutes)", "Estimated time to complete; shown to learners and printed on the certificate."],
        ["Mode", "Delivery label, e.g. “Online · Self-paced”."],
        ["Free / Price / Compare-at", "Pricing fields (cents). Payment is schema-level; most courses are free today."],
        ["Status", "Draft (hidden), Published (live on /courses), Archived (retired)."],
        ["Featured", "Flags the course for promoted placement."],
        ["Skip-to-cert", "Lets experienced learners go straight to the quiz/certification."],
        ["Skills", "Skill tags a learner earns on completion (see §7 — these surface on caregiver profiles)."],
    ],
    widths=[2.1, 5.3],
)

h2("Curriculum")
table(
    ["Term", "Meaning / effect"],
    [
        ["Chapter", "A group of modules; appears as a section in the player’s left “path”."],
        ["Module", "A single lesson: rich text and/or a video, optional resources, optional reflection."],
        ["Resource", "A downloadable file or external link attached to a module."],
        ["“Pause & Notice” (revision question)", "A reflective check inside a module. Options can be marked “highlight” (recommended) with an explanation. It is NOT graded — it’s for noticing/learning. This is the “revision quiz” seen in the player."],
        ["Mark complete", "Learners mark each module done to advance; progress is tracked per enrollment."],
    ],
    widths=[2.3, 5.1],
)

h2("Quiz (the “Integration Moment”)")
table(
    ["Term", "Meaning / effect"],
    [
        ["Integration Moment", "The course’s final graded quiz."],
        ["Intro copy", "Short text shown before the quiz starts."],
        ["Pass threshold %", "Minimum score to pass and be certified (default 60%)."],
        ["Question / Options", "Multiple-choice. Exactly one option is marked correct."],
        ["Explanation", "Per-option rationale, revealed to the learner after they pass."],
        ["Attempts", "Unlimited; right/wrong is only revealed after passing."],
    ],
    widths=[2.1, 5.3],
)

# ---------------------------------------------------------------------------
h1("5. The learner experience (what you’re building for)")
numbered(" A learner finds a published course on /courses and opens the detail page.", bold_lead="Discover. ")
numbered(" Enrolling starts the course and unlocks the player.", bold_lead="Enroll. ")
numbered(" The player shows one module at a time with the chapter “path” on the right; the learner marks each module complete. “Pause & Notice” reflections appear inline.", bold_lead="Learn. ")
numbered(" After all modules, the learner takes the Integration Moment quiz and must hit the pass threshold (unlimited attempts).", bold_lead="Integrate. ")
numbered(" On passing (or completing, if there’s no quiz), a certificate is issued.", bold_lead="Certify. ")

# ---------------------------------------------------------------------------
h1("6. Certification")
bullet(" Completing the course (and passing the quiz, if any) issues a certificate to the learner.", bold_lead="When it’s issued. ")
bullet(" Each certificate has a unique Certificate ID and a verify token, and prints the course title, learner name, mode, estimated learning time, the Signers, and the footer disclaimer you set on the Certificate tab.", bold_lead="What’s on it. ")
bullet(" Anyone can confirm a certificate at /verify/[token] (valid / revoked / who / when). This is how families and programs trust a caregiver’s training.", bold_lead="Public verification. ")
bullet(" The certificate is viewable by the learner at /courses/[slug]/certificate.", bold_lead="Where the learner sees it. ")

# ---------------------------------------------------------------------------
h1("7. Skills & badges → caregiver profiles & the marketplace")
bullet(" Skills you attach to a course (Course details) are awarded to the learner on completion.")
bullet(" Earned, course-backed skills appear on the caregiver’s public profile (badged, e.g. “TRC Certified Pro”) and the course credential is listed under Education & Credentials.")
bullet(" This is the connection between authoring a good course and a caregiver’s credibility in “Find Caregivers”: completed courses become visible, verifiable training on their profile.")

# ---------------------------------------------------------------------------
h1("8. Admin Preview (new feature)")
para("A new admin-only Preview lets you walk a course exactly as a learner would — without enrolling and "
     "without writing any progress — so you can sanity-check content before publishing.", italic=True)
bullet(" The “Preview” link on each row of /admin/courses, and the “Preview” button in the course editor header. Both open in a new tab.", bold_lead="Where. ")
bullet(" The full learner-style view: chapter/module navigation, videos, rich text, resources, the “Pause & Notice” reflection (click an option to reveal the highlighted response), and the quiz as an answer key (every question with the correct option and explanations marked).", bold_lead="What you see. ")
bullet(" The route lives under /admin (guarded by the admin check) at /admin/courses/[id]/preview. Nothing is saved — no enrollment, no progress, no quiz attempts.", bold_lead="Admin-only & read-only. ")
bullet(" It works for Draft courses too (so you preview before publishing); an amber “Admin preview” banner shows the status and a link back to the editor.", bold_lead="Drafts. ")

# ---------------------------------------------------------------------------
h1("9. Notes & gotchas")
bullet(" Save often. “Save course” persists the entire tree (details, curriculum, quiz, certificate) at once; the bottom bar shows “Saved.”.")
bullet(" Draft vs Published. Only Published courses appear on /courses. Use Preview to review drafts.")
bullet(" The quiz’s correct answers/explanations are hidden from learners until they pass; the admin Preview shows them as an answer key.")
bullet(" Email is copy-link only today (SMTP not configured) — relevant for any invite-style flows, not course authoring itself.")

para()
para("Companion videos (in the project root): admin-course-creation-demo-*.webm (create a course + Preview). "
     "Related: caregiver-profile-edit-demo and caregiver-reviews-demo.", italic=True, color=MUTED, size=9)

# ---------------------------------------------------------------------------
out = os.path.abspath(os.path.join(os.getcwd(), "..", "Course-Creation-Handover.docx"))
doc.save(out)
print("Saved:", out)
